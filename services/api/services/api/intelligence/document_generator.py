"""
Document Generation Engine — YU-NA Intelligence / Terra.OS
Faza 4: ASSEMBLY

Zgodnie ze specyfikacją §43 + §52 Offer Engine Opus Magnum v1.0:
- Jinja2 template engine → python-docx (DOCX output)
- WeasyPrint dla kosztorysu PDF
- 6 dokumentów: Formularz Oferty + Zał.1-4 + Kosztorys PDF
- Smart contract selection (Zał.3): CPV match + value + recency
- 5-krokowa walidacja pól przed generowaniem
- Output: ZIP package gotowy do e-Zamówienia
- Font: Times New Roman 11pt, marginesy 2.5cm (standard PZP)
"""

from __future__ import annotations

import io
import os
import re
import zipfile
from dataclasses import dataclass, field
from datetime import date, datetime
from decimal import Decimal
from hashlib import sha256
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Environment, FileSystemLoader, StrictUndefined

TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "documents"

# ---------------------------------------------------------------------------
# Data models (wejście do generatora)
# ---------------------------------------------------------------------------

@dataclass
class TenderContext:
    """Dane przetargu wymagane do generowania dokumentów."""
    nr_sprawy: str
    tytul: str
    zamawiajacy_nazwa: str
    cpv_kody: list[str] = field(default_factory=list)
    wartosc_brutto: Optional[Decimal] = None
    # Warunki udziału (wyekstrahowane z SWZ)
    warunek_zdolnosc_techniczna: Optional[str] = None
    warunek_sytuacja_finansowa: Optional[str] = None
    warunek_uprawnienia: Optional[str] = None
    wykluczenie_art109: Optional[str] = None
    # Termin
    termin_skladania: Optional[datetime] = None


@dataclass
class CompanyContext:
    """Profil firmy-wykonawcy."""
    nazwa_pelna: str
    nip: str
    adres_ulica: str
    adres_nr_budynku: str
    adres_kod_pocztowy: str
    adres_miasto: str
    regon: Optional[str] = None
    krs: Optional[str] = None
    referencje: list[dict] = field(default_factory=list)
    osoby_kluczowe: list[dict] = field(default_factory=list)


@dataclass
class KosztorysContext:
    """Zatwierdzony kosztorys ofertowy."""
    total_netto: Decimal
    total_brutto: Decimal
    vat_stawka: Decimal
    vat_kwota: Decimal
    pozycje: list[dict] = field(default_factory=list)
    cennik_okres: str = "Q2/2026"
    cennik_region: str = "śląskie"


@dataclass
class BidStrategy:
    """Strategia ofertowania (output z bid optimizer)."""
    termin_realizacji_dni: int
    gwarancja_miesiecy: int = 60
    wadium_forma: Optional[str] = None
    wadium_kwota: Optional[Decimal] = None
    podwykonawcy: Optional[str] = None
    termin_zwiazania_dni: int = 30


@dataclass
class GeneratedDocument:
    doc_type: str
    filename: str
    content: bytes
    checksum: str
    size_bytes: int


@dataclass
class OfertaPackage:
    documents: list[GeneratedDocument]
    manifest: dict
    zip_content: bytes
    zip_checksum: str
    total_size_bytes: int
    doc_count: int
    generated_at: datetime


# ---------------------------------------------------------------------------
# Jinja2 filters
# ---------------------------------------------------------------------------

def _format_pln(value) -> str:
    """Format: 927 149,00"""
    try:
        v = float(value)
        return f"{v:,.2f}".replace(",", " ").replace(".", ",")
    except (TypeError, ValueError):
        return str(value)


_JEDNOSCI = ["", "jeden", "dwa", "trzy", "cztery", "pięć", "sześć", "siedem",
             "osiem", "dziewięć"]
_NASTKI = ["", "jedenaście", "dwanaście", "trzynaście", "czternaście",
           "piętnaście", "szesnaście", "siedemnaście", "osiemnaście",
           "dziewiętnaście"]
_DZIESIATKI = ["", "dziesięć", "dwadzieścia", "trzydzieści", "czterdzieści",
               "pięćdziesiąt", "sześćdziesiąt", "siedemdziesiąt",
               "osiemdziesiąt", "dziewięćdziesiąt"]
_SETKI = ["", "sto", "dwieście", "trzysta", "czterysta", "pięćset",
          "sześćset", "siedemset", "osiemset", "dziewięćset"]


def _trzy_cyfry(n: int) -> str:
    if n == 0:
        return ""
    s = _SETKI[n // 100]
    reszta = n % 100
    if 11 <= reszta <= 19:
        s += (" " if s else "") + _NASTKI[reszta - 10]
    else:
        dz = _DZIESIATKI[reszta // 10]
        jd = _JEDNOSCI[reszta % 10]
        if dz:
            s += (" " if s else "") + dz
        if jd:
            s += (" " if s else "") + jd
    return s.strip()


def _kwota_slownie(value) -> str:
    """Zamień liczbę PLN na słownie: 'dziewięćset dwadzieścia siedem tysięcy...'"""
    try:
        v = round(float(value), 2)
    except (TypeError, ValueError):
        return str(value)

    grosze = round((v % 1) * 100)
    cale = int(v)

    mld = cale // 1_000_000_000
    cale %= 1_000_000_000
    mln = cale // 1_000_000
    cale %= 1_000_000
    tys = cale // 1_000
    jed = cale % 1_000

    parts = []
    if mld:
        parts.append(f"{_trzy_cyfry(mld)} miliardów")
    if mln:
        parts.append(f"{_trzy_cyfry(mln)} milionów")
    if tys:
        parts.append(f"{_trzy_cyfry(tys)} tysięcy")
    if jed:
        parts.append(_trzy_cyfry(jed))

    zlote = " ".join(parts) if parts else "zero"
    return f"{zlote} złotych {grosze:02d}/100"


# ---------------------------------------------------------------------------
# Jinja2 environment
# ---------------------------------------------------------------------------

def _build_jinja_env() -> Environment:
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATES_DIR)),
        undefined=StrictUndefined,
        trim_blocks=True,
        lstrip_blocks=True,
    )
    env.filters["format_pln"] = _format_pln
    env.filters["kwota_slownie"] = _kwota_slownie
    return env


# ---------------------------------------------------------------------------
# Smart reference selection (Załącznik 3)
# §52: CPV match + value match + recency
# ---------------------------------------------------------------------------

def _select_best_references(
    referencje: list[dict],
    tender_cpv: list[str],
    tender_wartosc: Optional[Decimal],
    required_count: int = 3,
) -> list[dict]:
    """
    Wybiera najlepsze referencje z profilu firmy.
    Scoring: CPV match (40) + value match (30) + recency (30).
    """
    if not referencje:
        return []

    scored = []
    for ref in referencje:
        score = 0.0

        # CPV match (40 pkt)
        ref_cpv = ref.get("cpv_kody", [])
        if ref_cpv and tender_cpv:
            overlap = len(set(ref_cpv) & set(tender_cpv))
            score += 40 * (overlap / max(len(tender_cpv), 1))
        elif not ref_cpv:
            score += 15  # brak CPV — częściowy bonus

        # Value match (30 pkt) — ref powinna być ≥ 50% wartości przetargu
        if tender_wartosc and ref.get("wartosc_brutto"):
            ratio = float(ref["wartosc_brutto"]) / float(tender_wartosc)
            if ratio >= 0.5:
                score += 30
            elif ratio >= 0.25:
                score += 15

        # Recency (30 pkt) — im nowsza, tym lepiej (max 5 lat wstecz)
        try:
            rok_ref = int(str(ref.get("data_zakonczenia", ""))[:4])
            rok_teraz = date.today().year
            wiek = rok_teraz - rok_ref
            score += max(0, 30 - wiek * 6)
        except (ValueError, TypeError):
            score += 10

        scored.append((score, ref))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [r for _, r in scored[:required_count]]


# ---------------------------------------------------------------------------
# Document generators
# ---------------------------------------------------------------------------

def _render_template(env: Environment, template_name: str, ctx: dict) -> str:
    """Renderuj szablon Jinja2 → tekst."""
    tmpl = env.get_template(template_name)
    return tmpl.render(**ctx)


def _text_to_docx(text: str, title: str) -> bytes:
    """
    Konwertuj plain text → DOCX.
    Font: Times New Roman 11pt, marginesy 2.5cm (standard PZP).
    """
    doc = Document()

    # Marginesy 2.5cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styl domyślny: Times New Roman 11pt
    from docx.styles.style import _CharacterStyle  # noqa: PLC0415
    style = doc.styles["Normal"]
    if hasattr(style, "font"):
        style.font.name = "Times New Roman"  # type: ignore[union-attr]
        style.font.size = Pt(11)  # type: ignore[union-attr]

    # Tytuł
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_paragraph("")

    # Treść linijka po linijce
    for line in text.split("\n"):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _kosztorys_to_docx(kosztorys: KosztorysContext, tender: TenderContext,
                       company: CompanyContext) -> bytes:
    """
    Generuj kosztorys ofertowy jako DOCX z tabelą pozycji.
    Format: zgodny z ATH/Norma (kolumny: Lp, Opis, KNR, JM, Ilość, Cena JM, Wartość).
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    if hasattr(style, "font"):
        style.font.name = "Times New Roman"  # type: ignore[union-attr]
        style.font.size = Pt(10)  # type: ignore[union-attr]

    # Nagłówek
    h = doc.add_heading("KOSZTORYS OFERTOWY", 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Postępowanie: {tender.tytul}")
    doc.add_paragraph(f"Nr sprawy: {tender.nr_sprawy}")
    doc.add_paragraph(f"Wykonawca: {company.nazwa_pelna} (NIP: {company.nip})")
    doc.add_paragraph(f"Ceny wg: Sekocenbud {kosztorys.cennik_okres}, region: {kosztorys.cennik_region}")
    doc.add_paragraph("")

    # Tabela pozycji
    headers = ["Lp", "Opis roboty", "KNR", "JM", "Ilość", "Cena JM netto [PLN]", "Wartość netto [PLN]"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Nagłówki
    hdr_row = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_row[i].text = h_text
        for para in hdr_row[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Pozycje
    for pozycja in kosztorys.pozycje:
        row = table.add_row().cells
        row[0].text = str(pozycja.get("lp", ""))
        row[1].text = pozycja.get("opis_roboty", "")
        knr = f"{pozycja.get('knr_katalog', '')} {pozycja.get('knr_tablica', '')}".strip()
        row[2].text = knr
        row[3].text = pozycja.get("jednostka", "")
        row[4].text = str(pozycja.get("ilosc", ""))
        row[5].text = _format_pln(pozycja.get("cena_jm_netto", 0))
        row[6].text = _format_pln(pozycja.get("wartosc_netto", 0))
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    # Podsumowanie
    doc.add_paragraph("")
    doc.add_paragraph(f"Wartość netto:   {_format_pln(kosztorys.total_netto)} PLN")
    doc.add_paragraph(f"VAT ({kosztorys.vat_stawka}%):       {_format_pln(kosztorys.vat_kwota)} PLN")
    p = doc.add_paragraph(f"WARTOŚĆ BRUTTO:  {_format_pln(kosztorys.total_brutto)} PLN")
    for run in p.runs:
        run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Validation (§52 — 5-step)
# ---------------------------------------------------------------------------

class DocumentValidationError(Exception):
    def __init__(self, errors: list[str]):
        self.errors = errors
        super().__init__(f"Walidacja dokumentów nieudana: {errors}")


def _validate_inputs(
    tender: TenderContext,
    company: CompanyContext,
    kosztorys: KosztorysContext,
    bid: BidStrategy,
) -> list[str]:
    """
    5-krokowa walidacja wejść przed generowaniem.
    Returns: lista błędów (pusta = OK).
    """
    errors = []

    # Step 1: Pola obowiązkowe firmy
    for field_name, label in [
        ("nazwa_pelna", "Nazwa firmy"),
        ("nip", "NIP"),
        ("adres_miasto", "Miasto"),
    ]:
        if not getattr(company, field_name, None):
            errors.append(f"[FIRMA] Brakuje: {label}")

    # Step 2: Walidacja NIP
    nip = re.sub(r"[^0-9]", "", company.nip or "")
    if len(nip) != 10:
        errors.append(f"[FIRMA] NIP nieprawidłowy: '{company.nip}' (musi mieć 10 cyfr)")

    # Step 3: Kwoty logiczne
    if kosztorys.total_brutto <= 0:
        errors.append("[KOSZTORYS] Wartość brutto musi być > 0")
    if kosztorys.total_netto > kosztorys.total_brutto:
        errors.append("[KOSZTORYS] Wartość netto > brutto — błąd VAT")

    # Step 4: Termin realizacji
    if bid.termin_realizacji_dni <= 0:
        errors.append("[OFERTA] Termin realizacji musi być > 0 dni")

    # Step 5: Gwarancja
    if not (12 <= bid.gwarancja_miesiecy <= 120):
        errors.append(f"[OFERTA] Gwarancja {bid.gwarancja_miesiecy} mies. poza zakresem 12–120")

    return errors


# ---------------------------------------------------------------------------
# Main orchestrator: DocumentOrchestrator
# ---------------------------------------------------------------------------

class DocumentOrchestrator:
    """
    Główny orchestrator generowania dokumentów ofertowych.
    Zgodnie z §43 pipeline: Formularz → Zał.1-4 → Kosztorys → PackageAssembler.
    """

    def __init__(self):
        self.jinja = _build_jinja_env()

    def generate_package(
        self,
        tender: TenderContext,
        company: CompanyContext,
        kosztorys: KosztorysContext,
        bid: BidStrategy,
    ) -> OfertaPackage:
        """
        Generuj kompletny pakiet ofertowy ZIP.
        Raises DocumentValidationError jeśli walidacja nie przejdzie.
        """
        # --- Walidacja ---
        errors = _validate_inputs(tender, company, kosztorys, bid)
        if errors:
            raise DocumentValidationError(errors)

        today_str = date.today().strftime("%d.%m.%Y")

        # Smart reference selection (§52)
        best_refs = _select_best_references(
            company.referencje,
            tender.cpv_kody,
            tender.wartosc_brutto,
        )

        # Kontekst bazowy Jinja2
        base_ctx: dict[str, Any] = {
            "tender": tender,
            "company": company,
            "cost": kosztorys,
            "bid": bid,
            "today": today_str,
            "referencje": best_refs,
            "osoby": company.osoby_kluczowe,
            "warunki": {
                "zdolnosc_techniczna": tender.warunek_zdolnosc_techniczna,
                "sytuacja_finansowa": tender.warunek_sytuacja_finansowa,
                "uprawnienia": tender.warunek_uprawnienia,
            },
        }

        documents: list[GeneratedDocument] = []

        # 1. Formularz Oferty
        documents.append(self._gen_doc(
            "formularz_oferty.jinja2",
            "FORMULARZ OFERTY",
            f"Formularz_Oferty_{tender.nr_sprawy.replace('/', '_')}.docx",
            "formularz_oferty",
            base_ctx,
        ))

        # 2. Załącznik 1 — Oświadczenie art. 108 (wykluczenia)
        documents.append(self._gen_doc(
            "zal1_oswiadczenie_art108.jinja2",
            "OŚWIADCZENIE art. 108 PZP",
            "Zal1_Oswiadczenie_art108_PZP.docx",
            "oswiadczenie_art108",
            base_ctx,
        ))

        # 3. Załącznik 2 — Oświadczenie art. 112 (warunki udziału)
        documents.append(self._gen_doc(
            "zal2_oswiadczenie_art112.jinja2",
            "OŚWIADCZENIE art. 112 PZP",
            "Zal2_Oswiadczenie_art112_PZP.docx",
            "oswiadczenie_art112",
            base_ctx,
        ))

        # 4. Załącznik 3 — Wykaz robót (smart selection)
        documents.append(self._gen_doc(
            "zal3_wykaz_robot.jinja2",
            "WYKAZ ROBÓT BUDOWLANYCH",
            "Zal3_Wykaz_Robot.docx",
            "wykaz_robot",
            base_ctx,
        ))

        # 5. Załącznik 4 — Wykaz osób
        documents.append(self._gen_doc(
            "zal4_wykaz_osob.jinja2",
            "WYKAZ OSÓB",
            "Zal4_Wykaz_Osob.docx",
            "wykaz_osob",
            base_ctx,
        ))

        # 6. Kosztorys ofertowy (DOCX z tabelą)
        kosztorys_bytes = _kosztorys_to_docx(kosztorys, tender, company)
        kosztorys_doc = self._make_doc(
            f"Kosztorys_Ofertowy_{tender.nr_sprawy.replace('/', '_')}.docx",
            "kosztorys_ofertowy",
            kosztorys_bytes,
        )
        documents.append(kosztorys_doc)

        # --- PackageAssembler: ZIP ---
        return self._assemble_zip(documents, tender)

    def _gen_doc(
        self,
        template: str,
        title: str,
        filename: str,
        doc_type: str,
        ctx: dict,
    ) -> GeneratedDocument:
        text = _render_template(self.jinja, template, ctx)
        docx_bytes = _text_to_docx(text, title)
        return self._make_doc(filename, doc_type, docx_bytes)

    @staticmethod
    def _make_doc(filename: str, doc_type: str, content: bytes) -> GeneratedDocument:
        checksum = sha256(content).hexdigest()
        return GeneratedDocument(
            doc_type=doc_type,
            filename=filename,
            content=content,
            checksum=checksum,
            size_bytes=len(content),
        )

    @staticmethod
    def _assemble_zip(
        documents: list[GeneratedDocument],
        tender: TenderContext,
    ) -> OfertaPackage:
        """
        PackageAssembler: pakuj dokumenty w ZIP z manifest.json.
        Zgodnie z §48 OfertaPackage schema.
        """
        manifest = {
            "created_at": datetime.utcnow().isoformat() + "Z",
            "tender_ref": tender.nr_sprawy,
            "documents": [
                {
                    "id": i + 1,
                    "doc_type": doc.doc_type,
                    "filename": doc.filename,
                    "size_bytes": doc.size_bytes,
                    "checksum": f"sha256:{doc.checksum}",
                    "required": True,
                }
                for i, doc in enumerate(documents)
            ],
            "total_docs": len(documents),
            "validation_passed": True,
        }

        import json
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for doc in documents:
                zf.writestr(doc.filename, doc.content)
            zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))

        zip_bytes = zip_buf.getvalue()
        zip_checksum = sha256(zip_bytes).hexdigest()

        return OfertaPackage(
            documents=documents,
            manifest=manifest,
            zip_content=zip_bytes,
            zip_checksum=zip_checksum,
            total_size_bytes=len(zip_bytes),
            doc_count=len(documents),
            generated_at=datetime.utcnow(),
        )


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------

def generate_oferta_package(
    tender: TenderContext,
    company: CompanyContext,
    kosztorys: KosztorysContext,
    bid: BidStrategy,
) -> OfertaPackage:
    """
    Entry point dla routera FastAPI.
    Raises DocumentValidationError przy błędach walidacji.
    """
    orchestrator = DocumentOrchestrator()
    return orchestrator.generate_package(tender, company, kosztorys, bid)
